"""Generate a concise PM reference brief (.docx) from the BTC quality report.

Aggregate stats only (NDA-safe: no real product records). Bullet-point style,
Vietnamese, one heading per requested section.

Usage: python scripts/gen_pm_brief.py
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "artifacts" / "btc_quality_report.json"
OUT = ROOT / "DMX_AI_Advisor_PM_Brief.docx"

GREEN = RGBColor(0x1B, 0x7A, 0x3D)


def main() -> int:
    rep = json.loads(REPORT.read_text(encoding="utf-8"))
    tot = rep["totals"]
    sheets = rep["by_sheet"]
    rows = []
    for info in sheets.values():
        u = info["unique_sku"]
        p = info["with_effective_price"]
        rows.append((info["category"], u, p, round(100 * p / max(u, 1)), info["eligible"]))
    rows.sort(key=lambda x: -x[1])
    sum_sku = sum(r[1] for r in rows)
    sum_price = sum(r[2] for r in rows)
    price_pct = round(100 * sum_price / sum_sku)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h1(text):
        h = doc.add_heading(text, level=1)
        h.runs[0].font.color.rgb = GREEN

    def h2(text):
        doc.add_heading(text, level=2)

    def b(text, level=0):
        p = doc.add_paragraph(style="List Bullet")
        if level:
            p.paragraph_format.left_indent = Pt(18 * (level + 1))
        p.add_run(text)

    # --- title ---
    title = doc.add_heading("Điện Máy Xanh — AI Product Advisor", level=0)
    title.runs[0].font.color.rgb = GREEN
    sub = doc.add_paragraph()
    r = sub.add_run("Bản tóm tắt cho Product Management — dựa trên phân tích dataset thô "
                    "(Spec_cate_gia.xlsx)")
    r.italic = True
    doc.add_paragraph("VAIC 2026 · Track Năng suất SME · Chatbot tư vấn mua hàng (chỉ text).")

    # --- 1. Key insights ---
    h1("1. Key Insights (từ dữ liệu thật)")
    b(f"Dataset ban tổ chức: {len(sheets)} ngành hàng, {tot['input_rows']:,} dòng, "
      f"{sum_sku:,} SKU unique (SKU = khóa chính, 1 model có nhiều SKU).")
    b(f"Dữ liệu THƯA là đặc điểm lớn nhất: chỉ ~{price_pct}% SKU có giá "
      f"({sum_price:,}/{sum_sku:,}). Nhiều field null/để trống theo nguồn — không phải lỗi xử lý.")
    b("Không có tồn kho, tên sản phẩm, ảnh hay URL trong file → không thể suy đoán, "
      "phải chờ API/feed bổ sung. Guardrail 'chưa có dữ liệu' là bắt buộc, không phải tùy chọn.")
    b("Độ phủ giá lệch mạnh theo ngành: cao ở smartwatch/water_heater/freezer/phone_mic "
      "(44–73%), thấp ở monitor/washing_machine/refrigerator (14–15%).")
    b("Máy lạnh là ngành demo tốt nhất: có phạm vi m², BTU, độ ồn, nhãn năng lượng, inverter "
      "— đủ tín hiệu để lọc + rank theo nhu cầu thật. Đã chuẩn hóa sâu 10 trường.")
    b(f"Sau làm sạch: {tot['eligible']:,} SKU đủ điều kiện demo (có giá + trường lõi); "
      f"riêng máy lạnh 152 SKU đủ đầy đủ tiêu chí (đời ≥2025, giá, diện tích, độ ồn, inverter).")
    # per-category mini table
    h2("Độ phủ giá theo ngành (SKU / có giá / %)")
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, txt in enumerate(["Ngành hàng", "SKU", "Có giá", "%"]):
        hdr[i].paragraphs[0].add_run(txt).bold = True
    for cat, u, p, pct, _ in rows:
        c = t.add_row().cells
        c[0].text, c[1].text, c[2].text, c[3].text = cat, f"{u:,}", f"{p:,}", f"{pct}%"

    # --- 2. Primary users ---
    h1("2. Primary Users")
    b("Khách mua sắm online phổ thông: mô tả nhu cầu bằng tiếng Việt tự nhiên "
      "(có lỗi chính tả, viết tắt, đơn vị m²/BTU/lít), không rành thông số kỹ thuật.")
    b("Nhân viên tư vấn bán hàng: dùng chatbot như trợ lý để trả lời nhanh, giảm câu hỏi lặp.")
    b("Đội vận hành thương mại điện tử / bán lẻ nhiều SKU: cần tăng tỷ lệ chuyển đổi traffic → đơn.")
    b("Người phụ trách chọn mua hộ (con mua cho bố mẹ, vợ/chồng): cần giải thích trade-off dễ hiểu.")

    # --- 3. Value proposition ---
    h1("3. Value Proposition")
    b("Hiểu NHU CẦU THẬT, không chỉ liệt kê thông số: hỏi ngược khi thiếu (phòng ngủ/khách, "
      "nắng, diện tích, ưu tiên êm hay lạnh nhanh, trả góp/khuyến mãi).")
    b("So sánh bằng ngôn ngữ bình dân, tập trung lợi ích thực tế — không thuật ngữ marketing.")
    b("Top 3 sản phẩm kèm lý do + đánh đổi + nguồn dữ liệu; nói rõ sản phẩm nào KHÔNG nên chọn.")
    b("Chống bịa tuyệt đối: giá/thông số/tồn/khuyến mãi đều có nguồn; thiếu thì nói 'chưa có dữ liệu'.")
    b("Giá trị thương mại: giảm tải tư vấn giờ cao điểm, cá nhân hóa, tăng conversion trên kênh online.")

    # --- 4. UX flow ---
    h1("4. UX Flow")
    b("B1 — Khách nhập nhu cầu tự nhiên: 'máy lạnh dưới 20 triệu, phòng 18m², tiết kiệm điện, ít ồn'.")
    b("B2 — LLM trích xuất slot: ngân sách, diện tích, ưu tiên, ràng buộc (< 3 giây).")
    b("B3 — Code kiểm tra slot thiếu → hỏi ngược đúng trọng tâm (phòng ngủ/khách? có nắng? trả góp?).")
    b("B4 — Code lọc hard-constraint + rank (thuần code, in-memory, vài ms) → không bịa số, luôn < 5 giây.")
    b("B5 — LLM giải thích Top 3 CHỈ từ dữ liệu đã chọn: ưu/nhược, vì sao hợp, nguồn (< 5 giây).")
    b("B6 — Thiếu giá/tồn → hiển thị 'chưa có dữ liệu', không đoán; khách quyết định nhanh hơn.")

    # --- 5. Feasibility ---
    h1("5. Tính khả thi")
    b("Kỹ thuật: pipeline dữ liệu đã xong (Phase 1) — schema 14 ngành, làm sạch, report chất lượng.")
    b("Latency: ranking/lọc = thuần code (~ms) → thừa sức < 5s; nút thắt duy nhất là inference LLM.")
    b("Chi phí: $30 credit FPT AI Factory ≈ ~70.000 lượt hội thoại → KHÔNG phải giới hạn "
      "(pilot chỉ cần 1.000–10.000). Test mock LLM = $0.")
    b("On-premise: dùng FPT AI Factory (host tại VN, OpenAI-compatible) — hợp yêu cầu, "
      "tránh phụ thuộc API nước ngoài đắt/bất ổn.")
    b("Rủi ro chính = DỮ LIỆU THƯA: chỉ ~25% SKU có giá. Giảm thiểu bằng cách demo trên tập "
      "eligible + guardrail trung thực; mở rộng khi ban tổ chức cấp data đầy đủ.")

    # --- 6. Pilot pathway ---
    h1("6. Pilot Pathway")
    b("Phạm vi: bắt đầu 1 ngành hàng — MÁY LẠNH (dữ liệu đặc trưng đầy đủ nhất để rank theo nhu cầu).")
    b("Thời gian: 3 tháng sau hackathon, khi có catalog/giá/tồn/chính sách mẫu đầy đủ.")
    b("Quy mô: 1.000–10.000 lượt hội thoại thử nghiệm trên 1 website/app bán lẻ.")
    b("Điều kiện ký hợp đồng: đạt KPI độ đúng dữ liệu, không hallucination nghiêm trọng, "
      "UX dễ dùng, có log nguồn, tích hợp được API catalog/stock/promotion.")
    b("Mở rộng dần: máy lạnh → tủ lạnh/máy giặt/điện thoại khi độ phủ dữ liệu tăng.")

    # --- 7. Pitch storyline ---
    h1("7. Pitch Storyline")
    b("Mở: 'Khách vào web, gõ một câu tiếng Việt đời thường — và được tư vấn như gặp nhân viên giỏi.'")
    b("Đau: bảng so sánh thông số khô khan → khách phổ thông không hiểu, quyết định chậm, traffic không convert.")
    b("Giải: trợ lý AI hỏi ngược đúng chỗ → lọc/rank bằng code → giải thích trade-off bình dân.")
    b("Khác biệt: KHÔNG bịa. Mọi con số có nguồn; thiếu thì nói thẳng 'chưa có dữ liệu'.")
    b("Chứng minh: demo thật trên máy lạnh — Top 3 kèm lý do, < 3s hỏi/đáp, < 5s so sánh.")
    b("Đóng: $30 đủ ~70k hội thoại, on-prem VN, sẵn sàng pilot 3 tháng → tăng conversion, giảm tải tư vấn.")

    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run("Nguồn: artifacts/btc_quality_report.json (thống kê tổng hợp, không chứa "
                      "bản ghi sản phẩm thật — an toàn NDA).")
    fr.italic = True
    fr.font.size = Pt(9)

    doc.save(OUT)
    print(f"saved {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
